from docx import Document
from pathlib import Path
import json
import re


DOCX_PATH = Path("thesis_review_source.docx")


def norm(text: str) -> str:
    return " ".join((text or "").split())


def is_heading(style: str, text: str) -> bool:
    if "Heading" in style or "标题" in style:
        return True
    return bool(
        re.match(r"^(第[一二三四五六七八九十]+章|\d+(?:\.\d+){0,3}\s|[一二三四五六七八九十]+、)", text)
    )


def is_chapter(text: str) -> bool:
    return bool(re.match(r"^第[0-9一二三四五六七八九十]+章", text) or re.match(r"^\d+\s+[^.\d]", text))


doc = Document(str(DOCX_PATH))

paras = []
for i, para in enumerate(doc.paragraphs):
    text = norm(para.text)
    if not text:
        continue
    style = para.style.name if para.style else ""
    paras.append({"i": i, "style": style, "text": text})

heads = []
for idx, item in enumerate(paras):
    if is_heading(item["style"], item["text"]):
        heads.append((idx, item["i"], item["style"], item["text"]))

chapter_indices = []
for idx, item in enumerate(paras):
    if is_chapter(item["text"]):
        chapter_indices.append((idx, item["text"]))

chapters = []
for n, (start, title) in enumerate(chapter_indices):
    end = chapter_indices[n + 1][0] if n + 1 < len(chapter_indices) else len(paras)
    text = "\n".join(item["text"] for item in paras[start:end])
    chapters.append(
        {
            "title": title,
            "paragraphs": end - start,
            "chars": len(text),
            "subheadings": [h[3] for h in heads if start <= h[0] < end and h[0] != start],
        }
    )

keywords = [
    "实现",
    "功能",
    "模块",
    "数据库",
    "界面",
    "测试",
    "需求",
    "设计",
    "代码",
    "Activity",
    "SQLite",
    "Android",
    "RecyclerView",
    "登录",
    "日程",
    "提醒",
]

print("PARAGRAPHS", len(paras))
print("TABLES", len(doc.tables))
print("INLINE_SHAPES", len(doc.inline_shapes))
print("\nHEADINGS:")
for h in heads[:300]:
    print(f"{h[1]}\t{h[2]}\t{h[3]}")

print("\nCHAPTERS:")
for chapter in chapters:
    print(json.dumps(chapter, ensure_ascii=False))

print("\nOPENING_SAMPLE:")
for item in paras[:45]:
    print(f"[{item['i']}] {item['style']}: {item['text'][:180]}")

print("\nKEY_IMPLEMENTATION_PARAS:")
for item in paras:
    if any(k in item["text"] for k in keywords) and len(item["text"]) > 30:
        print(f"[{item['i']}] {item['style']}: {item['text'][:260]}")

print("\nSECTION_WINDOWS:")
for start, end in [(113, 148), (148, 175), (175, 227), (228, 250)]:
    print(f"\n--- paragraphs {start}-{end} ---")
    for item in paras:
        if start <= item["i"] <= end:
            print(f"[{item['i']}] {item['style']}: {item['text'][:500]}")

print("\nTABLES_CONTENT:")
for ti, table in enumerate(doc.tables, start=1):
    print(f"\nTABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows[:12], start=1):
        cells = [norm(cell.text)[:120] for cell in row.cells]
        print(f"  R{ri}: {' | '.join(cells)}")
